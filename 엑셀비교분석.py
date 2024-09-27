import gradio as gr
import pandas as pd
import json
import difflib
from collections import OrderedDict
import openai
import os
import tempfile
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime

# 환경 변수 로드
load_dotenv()
openai_api_key = os.getenv('OPENAI_API_KEY')

if not openai_api_key:
    raise ValueError("OpenAI API 키가 설정되지 않았습니다. .env 파일에 'OPENAI_API_KEY'를 추가해주세요.")
else:
    openai.api_key = openai_api_key

class ExcelMarkdownJSONConverter:
    def __init__(self, data_dir='./data', json_file_path='./output_report.json'):
        """
        클래스 초기화.
        데이터 디렉토리에서 엑셀 파일을 로드하여 JSON 파일을 생성하고, JSON 데이터를 로드합니다.

        :param data_dir: 엑셀 파일이 저장된 디렉토리 경로
        :param json_file_path: JSON 파일을 저장할 경로
        """
        self.data_dir = data_dir
        self.json_file_path = json_file_path
        self.json_data = []

        # JSON 파일 생성 또는 로드
        if not os.path.isfile(self.json_file_path):
            print(f"[Info] JSON 파일을 생성합니다: {self.json_file_path}")
            self.generate_json_from_excel()
        else:
            print(f"[Info] JSON 파일을 로드합니다: {self.json_file_path}")
            self.load_json()

    def generate_json_from_excel(self):
        """
        데이터 디렉토리 내의 모든 엑셀 파일을 읽고 JSON 파일을 생성하여 저장합니다.
        """
        all_data = []
        excel_files = [file for file in os.listdir(self.data_dir) if file.endswith(('.xlsx', '.xls'))]
        
        if not excel_files:
            print(f"[Warning] '{self.data_dir}' 디렉토리에 엑셀 파일이 존재하지 않습니다.")
            return

        for excel_file in excel_files:
            file_path = os.path.join(self.data_dir, excel_file)
            try:
                excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
                for sheet_name, content in excel_data.items():
                    # 시트 내용을 문자열로 변환 (필요에 따라 조정 가능)
                    content_str = content.to_markdown(index=False)
                    all_data.append({
                        "file": excel_file,
                        "sheet_name": sheet_name,
                        "content": content_str
                    })
                print(f"[Success] '{excel_file}' 파일에서 시트를 성공적으로 추출했습니다.")
            except Exception as e:
                print(f"[Error] '{excel_file}' 파일을 처리하는 중 오류가 발생했습니다: {e}")

        # JSON 파일로 저장
        try:
            with open(self.json_file_path, 'w', encoding='utf-8') as json_file:
                json.dump(all_data, json_file, ensure_ascii=False, indent=4)
            self.json_data = all_data
            print(f"[Success] JSON 파일이 '{self.json_file_path}' 경로에 성공적으로 저장되었습니다.")
        except Exception as e:
            print(f"[Error] JSON 파일을 저장하는 중 오류가 발생했습니다: {e}")

    def load_json(self):
        """
        JSON 파일을 로드하여 클래스 내부에 저장합니다.
        """
        try:
            with open(self.json_file_path, 'r', encoding='utf-8') as json_file:
                self.json_data = json.load(json_file)
            print(f"[Success] JSON 파일이 성공적으로 로드되었습니다.")
        except Exception as e:
            print(f"[Error] JSON 파일을 로드하는 중 오류가 발생했습니다: {e}")

    def search_similar_sheets(self, selected_sheet_name, top_n=3, min_similarity=0.5):
        """
        선택된 시트 이름과 유사한 시트들을 JSON 데이터에서 검색합니다.

        :param selected_sheet_name: 사용자가 선택한 시트 이름
        :param top_n: 반환할 최대 유사 시트의 수
        :param min_similarity: 결과로 반환할 최소 유사도 점수
        :return: 유사한 시트들의 리스트
        """
        # 중복 방지를 위해 고유한 시트 이름 추출
        unique_sheets = list(OrderedDict.fromkeys([item['sheet_name'] for item in self.json_data]))
        best_matches = self.find_most_similar_sheets(selected_sheet_name, unique_sheets, top_n, min_similarity)

        results = []
        for match, similarity in best_matches:
            # 매칭된 시트 이름을 가진 모든 시트 찾기
            matched_sheets = [sheet for sheet in self.json_data if sheet['sheet_name'] == match]
            for sheet in matched_sheets:
                results.append({
                    "입력": selected_sheet_name,
                    "파일명": sheet.get('file', '알 수 없음'),
                    "가장 유사한 시트 이름": sheet['sheet_name'],
                    "유사도 점수": round(similarity, 2),
                    "Markdown 내용": sheet['content']
                })

        # 최소 유사도 점수 이상인 결과만 필터링
        filtered_results = [result for result in results if result['유사도 점수'] >= min_similarity]

        return filtered_results

    @staticmethod
    def find_most_similar_sheets(user_input, sheet_names, top_n=3, min_similarity=0.5):
        """
        사용자의 입력과 가장 유사한 시트 이름들을 찾습니다.

        :param user_input: 사용자가 입력한 문자열
        :param sheet_names: 시트 이름 리스트 (unique)
        :param top_n: 반환할 최대 유사 시트의 수
        :param min_similarity: 최소 유사도 점수
        :return: 유사한 시트 이름들과 유사도 점수의 리스트
        """
        # 대소문자 구분 없이 비교를 위해 소문자로 변환
        user_input_lower = user_input.lower()
        sheet_names_lower = [name.lower() for name in sheet_names]

        # difflib을 사용하여 가장 유사한 시트 이름 찾기
        matches = difflib.get_close_matches(user_input_lower, sheet_names_lower, n=top_n, cutoff=0.0)
        similar_sheets = []
        added_names = set()

        for match_lower in matches:
            # 동일한 시트를 여러 번 추가하지 않도록 방지
            if match_lower in added_names:
                continue
            try:
                index = sheet_names_lower.index(match_lower)
                original_name = sheet_names[index]
                similarity = difflib.SequenceMatcher(None, user_input_lower, match_lower).ratio()
                if similarity >= min_similarity:
                    similar_sheets.append((original_name, similarity))
                    added_names.add(match_lower)
            except ValueError:
                continue  # 매치가 없는 경우 무시

        return similar_sheets

    def generate_analysis_report(self, selected_sheet_name, similar_sheets, max_tokens=1500):
        """
        GPT를 사용하여 분석 보고서를 생성합니다.

        :param selected_sheet_name: 사용자가 선택한 시트 이름
        :param similar_sheets: 유사한 시트들의 리스트
        :param max_tokens: GPT 응답의 최대 토큰 수
        :return: 생성된 분석 보고서 문자열
        """
        # 선택된 시트의 내용을 가져오기
        selected_sheet_content = next((item['content'] for item in self.json_data if item['sheet_name'] == selected_sheet_name), "내용을 찾을 수 없습니다.")

        # 유사한 시트들의 내용 수집
        similar_contents = []
        for sheet in similar_sheets:
            similar_contents.append({
                "파일명": sheet['파일명'],
                "시트 이름": sheet['가장 유사한 시트 이름'],
                "유사도 점수": sheet['유사도 점수'],
                "내용": sheet['Markdown 내용']
            })

        # GPT에게 보낼 프롬프트 작성
        prompt = f"""
사용자가 선택한 시트 '{selected_sheet_name}'

**선택된 시트 내용:**
{selected_sheet_content}

**유사한 시트들:**
"""
        for idx, sheet in enumerate(similar_contents, start=1):
            prompt += f"""
{idx}. 파일명: {sheet['파일명']}
   시트 이름: {sheet['시트 이름']}
   유사도 점수: {sheet['유사도 점수']}
   내용:
   {sheet['내용']}
"""

        try:
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": """
다음 정보는 사용자의 엑셀 파일입니다. 이때 사용자의 엑셀 시트와 그와 유사한 시트들을 이애하고 특성을 분석해 주세요. 
차례대로 아래의 방법을 따릅니다. 

1. 사용자의 시트가 어떤 분야의 시트인지 파악하시오.
2. 사용자 시트의 주요 키워드나 핵심 요소는 무엇인지 파악하시오.
3. 사용자 시트의 키워드와 핵심 요소를 바탕으로 유사한 시트들과 사용자의 시트를 분석해주세요. 이때 키워드를 차례대로 비교 분석합니다. 
4. 또한 어떤 파일과 비교하여 어떤 수치를 기반으로 그렇게 판단하게 되었는지 자세하게 서술하시오. 
5. 수치 데이터를 비교할 경우, Table 로 정리해주세요. 

6. 요약 - 서론 - 본론 - 결론의 형태로 작성하시오. 
[결과물 형식]
요약 
서론 
본론 - 파일 분야 파악, 주요 키워드 및 핵심 분석 요소, 유사 자료와의 분석, 수치 기반 비교, 수치 기반 비교(상세 비교, 줄글의 형태) 
결과 
"""
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                max_tokens=3000,
                top_p=1
            )

            answer = response.choices[0].message.content
            return answer

        except Exception as e:
            print(f"[Error] GPT를 사용하여 분석 보고서를 생성하는 중 오류가 발생했습니다: {e}")
            return ""
            
    def generate_report2(self, report, max_tokens=1500):
        """
        GPT를 사용하여 보고서를 Markdown 형식으로 정리합니다.

        :param report: 기존 보고서 내용
        :param max_tokens: GPT 응답의 최대 토큰 수
        :return: 정리된 보고서 문자열
        """
        try:
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": """
당신은 유능한 데이터처리 전문가입니다. 아래의 데이터 분석본을 Markdown의 형식으로 보기쉽게 정리해주세요. 손실되는 정보는 없어야 합니다. 
아래의 순서대로 작업을 진행하시오. 

1. 글을 분석하고, 파트별로 분류합니다. 이때 글을 대표하는 제목에는 #를 붙여서 가장 크게보이도록 설정합니다. 
2. 파트에서 가장 중요한 부분에는 ## 을 달아서 강조해서 볼 수 있도록 합니다. 
3. 파트에서 가장 중요하지 않으나, 핵심적인 부분에는 ###를 달아서 일반 글과 다르게 보일 수 있도록 합니다. 
4. 각 파트를 하나로 모아서, 하나의 글이 될 수 있도록 하고.. 최종 완성본 앞에는 [완성본]을 달아서 완성본만을 볼 수 있도록 합니다. 
5. Table 부분은 Table로 정리해서 작성해주세요. Markdown 형식의 Table로 정리합니다. 
6. 각 파트를 구분할 수 있도록 --- 작성을 통해서 정렬해주세요. 
"""
                    },
                    {
                        "role": "user",
                        "content": report
                    }
                ],
                temperature=0.7,
                max_tokens=3000,
                top_p=1
            )

            answer = response.choices[0].message.content
            return answer
        except Exception as e:
            print(f"[Error] GPT를 사용하여 보고서를 정리하는 중 오류가 발생했습니다: {e}")
            return ""

    def parse_markdown_to_docx(self, markdown_content, query):
        """
        간단한 마크다운을 DOCX 파일으로 변환합니다.
        
        :param markdown_content: 마크다운 형식의 문자열
        :param query: 파일명에 포함될 쿼리 문자열
        :return: 저장된 파일의 전체 경로
        """
        # 파일명 생성
        output_filename = f'{query}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # 저장할 폴더 경로 설정 (현재 디렉토리)
        output_dir = "./"
        # 전체 파일 경로 생성
        output_path = os.path.join(output_dir, output_filename)
        
        # 워드 문서 생성
        doc = Document()
        
        # 스타일 설정
        styles = doc.styles

        # 제목 스타일 설정
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = '바탕체'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 부제목 스타일 설정
        if 'CustomHeading1' not in styles:
            heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = '바탕체'
            heading1_style.font.size = Pt(12)
            heading1_style.font.bold = True

        if 'CustomHeading2' not in styles:
            heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = '바탕체'
            heading2_style.font.size = Pt(10)
            heading2_style.font.bold = True

        # 본문 스타일 설정
        if 'CustomBody' not in styles:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = '바탕체'
            body_style.font.size = Pt(8)
            body_style.paragraph_format.line_spacing = 1.6
            body_style.font.color.rgb = RGBColor(99,99,99)

        # 리스트 스타일 설정 함수
        def add_bullet_paragraph(doc, text):
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # 테이블 데이터 임시 저장
        table_buffer = []
        in_table = False
        table_type = None  # 'pipe' or 'tab'

        # 마크다운 라인별 처리
        lines = markdown_content.split('\n')
        for idx, line in enumerate(lines):
            original_line = line  # 원본 라인 저장
            line = line.strip()
            if not line:
                continue

            # Pipe-separated table
            if re.match(r'^\|.*\|$', line):
                table_buffer.append(line)
                in_table = True
                table_type = 'pipe'
                continue
            # Tab-separated table
            elif line.count('\t') >= 1 and not in_table:
                table_buffer.append(line)
                in_table = True
                table_type = 'tab'
                continue
            elif in_table:
                if table_type == 'pipe' and re.match(r'^\|[-:| ]+\|$', line):
                    table_buffer.append(line)
                    continue
                elif table_type == 'tab' and re.match(r'^[-\t ]+$', line):
                    table_buffer.append(line)
                    continue
                else:
                    # 테이블 종료
                    if table_buffer:
                        # 테이블 파싱
                        table_data = []
                        for tbl_line in table_buffer:
                            if table_type == 'pipe':
                                row = [cell.strip() for cell in tbl_line.strip('|').split('|')]
                            elif table_type == 'tab':
                                row = [cell.strip() for cell in tbl_line.split('\t')]
                            table_data.append(row)
                        
                        # Remove separator lines
                        if table_type == 'pipe':
                            # Remove the second line (separator)
                            table_data = [row for i, row in enumerate(table_data) if i !=1]
                        elif table_type == 'tab':
                            # Assume second line is separator
                            table_data = [row for i, row in enumerate(table_data) if i !=1]
                        
                        # Find the maximum number of columns
                        max_cols = max(len(row) for row in table_data)
                        # Normalize all rows to have the same number of columns
                        for row in table_data:
                            while len(row) < max_cols:
                                row.append('')
                        
                        # 테이블 추가 (헤더 포함)
                        table = doc.add_table(rows=len(table_data), cols=max_cols)
                        table.style = 'Table Grid'  # 표 스타일 설정
                        
                        for row_idx, row_data in enumerate(table_data):
                            row = table.rows[row_idx]
                            for col_idx, cell in enumerate(row_data):
                                row.cells[col_idx].text = cell
                                if row_idx ==0:
                                    # 헤더 셀 굵게 처리
                                    for paragraph in row.cells[col_idx].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                        
                        # 테이블 버퍼 초기화
                        table_buffer = []
                        in_table = False
                        table_type = None

                    # 현재 라인은 테이블이 아니므로 계속 처리
                    # fall through to process the current line

            # 헤더 처리
            header_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if header_match:
                header_level = len(header_match.group(1))
                header_text = header_match.group(2)
                if header_level == 1:
                    doc.add_paragraph(header_text, style='CustomTitle')
                elif header_level == 2:
                    doc.add_paragraph(header_text, style='CustomHeading1')
                elif header_level == 3:
                    doc.add_paragraph(header_text, style='CustomHeading2')
                else:
                    doc.add_paragraph(header_text, style=f'Heading{header_level}')
                continue

            # 리스트 처리
            if line.startswith('- '):
                list_item = line[2:].strip()
                # 굵은 글씨 처리
                bold_match = re.findall(r'\*\*(.*?)\*\*', list_item)
                if bold_match:
                    # 리스트 항목에 굵은 글씨가 포함된 경우
                    p = doc.add_paragraph(style='List Bullet')
                    segments = re.split(r'(\*\*.*?\*\*)', list_item)
                    for segment in segments:
                        if segment.startswith('**') and segment.endswith('**'):
                            text = segment[2:-2]
                            run = p.add_run(f"{text} ")
                            run.bold = True
                        else:
                            p.add_run(segment)
                else:
                    add_bullet_paragraph(doc, list_item)
                continue

            # 일반 본문 처리
            # 굵은 글씨 처리
            bold_segments = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph(style='CustomBody')
            for segment in bold_segments:
                if segment.startswith('**') and segment.endswith('**'):
                    text = segment[2:-2]
                    run = p.add_run(text)
                    run.bold = True
                else:
                    p.add_run(segment)

        # 테이블 버퍼에 남아있는 데이터 처리 (파일 끝에 테이블이 있는 경우)
        if table_buffer:
            table_data = []
            for tbl_line in table_buffer:
                if table_type == 'pipe':
                    row = [cell.strip() for cell in tbl_line.strip('|').split('|')]
                elif table_type == 'tab':
                    row = [cell.strip() for cell in tbl_line.split('\t')]
                table_data.append(row)
            
            # Remove separator lines
            if table_type == 'pipe':
                table_data = [row for i, row in enumerate(table_data) if i !=1]
            elif table_type == 'tab':
                table_data = [row for i, row in enumerate(table_data) if i !=1]
            
            # Find the maximum number of columns
            max_cols = max(len(row) for row in table_data)
            # Normalize all rows to have the same number of columns
            for row in table_data:
                while len(row) < max_cols:
                    row.append('')
            
            # 테이블 추가 (헤더 포함)
            table = doc.add_table(rows=len(table_data), cols=max_cols)
            table.style = 'Table Grid'  # 표 스타일 설정
            
            for row_idx, row_data in enumerate(table_data):
                row = table.rows[row_idx]
                for col_idx, cell in enumerate(row_data):
                    row.cells[col_idx].text = cell
                    if row_idx ==0:
                        # 헤더 셀 굵게 처리
                        for paragraph in row.cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

        # DOCX 저장
        # doc.save(output_path)
        return output_path


# Gradio 애플리케이션 함수들
def upload_excel(excel_file):
    """
    사용자가 업로드한 엑셀 파일을 처리하고 시트 목록을 반환합니다.
    """
    if not excel_file:
        return gr.update(choices=[]), "엑셀 파일이 업로드되지 않았습니다."
    try:
        # 엑셀 파일이 파일 경로 문자열인 경우 직접 읽기
        if isinstance(excel_file, str):
            excel_path = excel_file
        else:
            # 파일 객체인 경우 임시 파일로 저장 후 경로 사용
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                tmp.write(excel_file.read())
                excel_path = tmp.name
        
        # 엑셀 파일을 읽어 시트 이름 추출
        excel_data = pd.read_excel(excel_path, sheet_name=None, engine='openpyxl')
        sheet_names = list(excel_data.keys())
        return gr.update(choices=sheet_names), "엑셀 파일이 성공적으로 로드되었습니다."
    except Exception as e:
        return gr.update(choices=[]), f"엑셀 파일을 로드하는 중 오류가 발생했습니다: {e}"

def search_similar(selected_sheet, converter: ExcelMarkdownJSONConverter):
    """
    선택된 시트와 유사한 시트를 검색합니다.
    """
    if not selected_sheet:
        return None, "시트가 선택되지 않았습니다."
    results = converter.search_similar_sheets(selected_sheet, top_n=3, min_similarity=0.5)
    if results:
        return results, "유사한 시트를 성공적으로 검색했습니다."
    else:
        return None, "유사한 시트를 찾을 수 없습니다."

def generate_report(selected_sheet, similar_sheets, converter: ExcelMarkdownJSONConverter):
    """
    GPT를 사용하여 분석 보고서를 생성합니다.
    """
    if not converter or not converter.json_data:
        return "", "JSON 데이터가 로드되지 않았습니다."
    report = converter.generate_analysis_report(selected_sheet, similar_sheets)
    report = converter.generate_report2(report)
    if report:
        return report, "분석 보고서가 성공적으로 생성되었습니다."
    else:
        return "", "분석 보고서 생성에 실패했습니다."

def download_report(report_text, converter: ExcelMarkdownJSONConverter, selected_sheet):
    """
    보고서를 DOCX 파일로 변환하여 다운로드할 수 있도록 합니다.
    """
    if not report_text:
        return None
    if not converter:
        return None
    query = selected_sheet if selected_sheet else "report"
    docx_path = converter.parse_markdown_to_docx(report_text, query)
    return docx_path

# Gradio 인터페이스 구성
with gr.Blocks() as demo:
    gr.Markdown("## Excel Markdown JSON Converter with GPT Analysis")

    # 초기화 및 JSON 생성
    converter = ExcelMarkdownJSONConverter()

    with gr.Row():
        with gr.Column():
            gr.Markdown("### 1. 엑셀 파일 업로드")
            excel_input = gr.File(label="엑셀 파일 업로드", file_types=['.xlsx', '.xls'])
            excel_status = gr.Textbox(label="엑셀 로드 상태", interactive=False)
            load_excel_button = gr.Button("엑셀 로드")
        with gr.Column():
            gr.Markdown("### 2. 시트 선택")
            sheet_dropdown = gr.Dropdown(label="시트 선택", choices=[], interactive=True)
            sheet_status = gr.Textbox(label="시트 선택 상태", interactive=False)

    with gr.Row():
        search_button = gr.Button("유사 시트 검색")
        search_status = gr.Textbox(label="검색 상태", interactive=False)

    with gr.Row():
        matched_sheets_output = gr.JSON(label="매칭된 시트들")

    with gr.Row():
        generate_report_button = gr.Button("분석 보고서 생성")
        report_status = gr.Textbox(label="보고서 생성 상태", interactive=False)

    with gr.Row():
        report_output = gr.Textbox(label="분석 보고서", lines=20)
        download_report_button = gr.Button("보고서 다운로드")
        download_report_file = gr.File(label="보고서 다운로드")

    # 상태 저장용 Gradio State
    converter_state = gr.State(converter)

    # 엑셀 파일 업로드 버튼 클릭 시 동작
    load_excel_button.click(
        upload_excel,
        inputs=excel_input,
        outputs=[sheet_dropdown, excel_status]
    )

    # 유사 시트 검색 버튼 클릭 시 동작
    search_button.click(
        search_similar,
        inputs=[sheet_dropdown, converter_state],
        outputs=[matched_sheets_output, search_status]
    )

    # 분석 보고서 생성 버튼 클릭 시 동작
    generate_report_button.click(
        generate_report,
        inputs=[sheet_dropdown, matched_sheets_output, converter_state],
        outputs=[report_output, report_status]
    )

    # 보고서 다운로드 버튼 클릭 시 동작
    download_report_button.click(
        download_report,
        inputs=[report_output, converter_state, sheet_dropdown],
        outputs=download_report_file
    )

demo.launch(debug=True, share=True)
